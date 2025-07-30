import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
import json
import os

# Configuração da página
st.set_page_config(
    page_title="Gerador de Laudos de Inspeção Predial",
    page_icon="🏢",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS customizado
st.markdown("""
    <style>
    .main { padding-top: 2rem; }
    .stButton>button {
        width: 100%;
        background-color: #1f77b4;
        color: white;
    }
    .stButton>button:hover {
        background-color: #145a8b;
    }
    </style>
""", unsafe_allow_html=True)

# Inicializar sessão
if 'laudos_salvos' not in st.session_state:
    st.session_state.laudos_salvos = {}
if 'eventos' not in st.session_state:
    st.session_state.eventos = []
if 'dados_laudo' not in st.session_state:
    st.session_state.dados_laudo = {}

# Dicionários de opções
OPCOES = {
    "contratada": [
        "Testcon Engenharia",
        "E2E Consultoria e Gestão",
        "Outra"
    ],
    "tipo_empreendimento": [
        "Institucional de ensino superior privado",
        "Comercial",
        "Residencial multifamiliar",
        "Industrial",
        "Hospitalar",
        "Outro"
    ],
    "anomalias": [
        "Eflorescência",
        "Pinturas em desconformidades",
        "Pilares apresentam expansão de armadura",
        "Marquises com rupturas e desplacamento",
        "Corrosão",
        "Mofo e bolor",
        "Infiltrações",
        "Fissuras",
        "Trincas",
        "Rachaduras",
        "Desplacamento de revestimento",
        "Vazamentos",
        "Problemas estruturais",
        "Deficiência de impermeabilização",
        "Instalações elétricas inadequadas",
        "Selantes inadequados",
        "Pintura deteriorada",
        "Desorganização",
        "Fixação inadequada",
        "Sem funcionamento",
        "Base de fixação inadequada",
        "Sistema inadequado",
        "Manchas de umidade",
        "Comprometimento de equipamentos",
        "Deficiência de ventilação",
        "Outra"
    ],
    "causas": [
        "Endógena",
        "Exógena", 
        "Funcional",
        "Endógena/Funcional",
        "Funcional/Exógena",
        "Outra"
    ],
    "consequencias": [
        "Prejuízo estético",
        "Iminência de infiltração",
        "Risco à segurança dos usuários",
        "Comprometimento estrutural",
        "Insalubridade",
        "Perda de funcionalidade",
        "Comprometimento de equipamentos",
        "Falta de acessibilidade",
        "Prejuízo estético e risco à segurança dos usuários",
        "Prejuízo estético, iminência de infiltração e risco à segurança dos usuários",
        "Prejuízo estético, insalubridade e risco à segurança dos usuários",
        "Outra"
    ],
    "recomendacoes": [
        "Contratar empresa especializada para reabilitar as estruturas",
        "Realizar pintura de toda área",
        "Revisar estruturas e trocar selantes",
        "Impermeabilizar áreas afetadas",
        "Adequar instalações elétricas",
        "Realizar limpeza e organização",
        "Substituir elementos danificados",
        "Realizar manutenção preventiva",
        "Contratar empresa para verificação e adequação",
        "Reabilitar pinturas das paredes e tetos",
        "Fazer limpeza na área",
        "Contratar empresa especializada para manutenção",
        "Contratar empresa para adequar circulação do ar",
        "Contratar empresa especializada para revisão de toda instalação elétrica",
        "Outra"
    ]
}

# Documentações padrão
DOCUMENTACOES = [
    "Certificado de Conclusão de Obra ou Habite-se",
    "Alvará ou Licença de Funcionamento",
    "Auto de Vistoria do Corpo de Bombeiros",
    "Licença de operação da ETE",
    "Licenças ambientais",
    "Certificado de Acessibilidade",
    "Licença de perfuração poços profundos",
    "Documentos de formação da brigada de incêndio",
    "Alvará de aprovação para instalação de equipamento",
    "Declaração de prestação de serviços de Pronto Atendimento",
    "Aprovação de paralelismo de Grupo Moto Gerador",
    "Manual de Uso, Operação e Manutenção",
    "Registros de manutenções",
    "Projetos Arquitetônicos"
]

# Funções auxiliares para gerar o documento
def gerar_documento_completo(dados, eventos, incluir_rodape=True, incluir_numeracao=True, versao=1):
    """Gera o documento Word completo"""
    doc = Document()
    
    # Configurar estilos básicos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # CAPA
    # Título principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RELATÓRIO DE ENGENHARIA")
    run.font.size = Pt(16)
    run.font.bold = True
    
    # Subtítulo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Laudo Técnico de Inspeção Predial")
    run.font.size = Pt(14)
    run.font.bold = True
    
    # Espaços
    for _ in range(3):
        doc.add_paragraph()
    
    # Informações do contratante
    p = doc.add_paragraph()
    p.add_run("Contratante: ").bold = True
    p.add_run(dados.get('contratante', ''))
    
    p = doc.add_paragraph()
    p.add_run("CNPJ: ").bold = True
    p.add_run(dados.get('cnpj', ''))
    
    p = doc.add_paragraph()
    p.add_run("Data: ").bold = True
    p.add_run(dados.get('data_laudo').strftime('%d/%m/%Y') if dados.get('data_laudo') else '')
    
    # Espaços
    for _ in range(5):
        doc.add_paragraph()
    
    # Imóvel
    p = doc.add_paragraph()
    p.add_run("Imóvel motivo:").bold = True
    
    p = doc.add_paragraph()
    p.add_run(dados.get('endereco', ''))
    
    # Espaços
    for _ in range(3):
        doc.add_paragraph()
    
    # Responsável técnico
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Silvio Augusto Barbosa de Albuquerque Filho, Engenheiro Civil")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("CREA/PE nº 054787D-PE")
    
    # Nova página para o sumário
    doc.add_page_break()
    
    # SUMÁRIO
    doc.add_heading('Sumário', level=1)
    
    # Lista de seções
    secoes = [
        ("1. RESSALVAS INICIAIS", "4"),
        ("2. OBJETIVO", "5"),
        ("3. DESCRIÇÃO DO OBJETO INSPECIONADO", "8"),
        ("4. REFERÊNCIAS NORMATIVAS", "11"),
        ("5. TERMINOLOGIA", "12"),
        ("6. ABRANGÊNCIA DA ANÁLISE", "18"),
        ("7. CLASSIFICAÇÃO DAS IRREGULARIDADES", "19"),
        ("8. PATAMARES DE CRITICIDADE", "20"),
        ("9. AVALIAÇÃO DE MANUTENÇÃO", "21"),
        ("10. AVALIAÇÃO DE USO", "23"),
        ("11. METODOLOGIA", "23"),
        ("12. DOCUMENTAÇÕES SOLICITADAS E DISPONIBILIZADAS", "26"),
        ("13. ANAMNESE", "27"),
        ("14. LAUDO TÉCNICO", "48"),
        ("15. DATA DO RELATÓRIO TÉCNICO", "53")
    ]
    
    for titulo, pagina in secoes:
        p = doc.add_paragraph()
        p.add_run(titulo)
        p.add_run(f" {'.'*50} {pagina}")
    
    # Nova página para o conteúdo
    doc.add_page_break()
    
    # 1. RESSALVAS INICIAIS
    doc.add_heading('RESSALVAS INICIAIS', level=1)
    doc.add_paragraph("O presente relatório técnico obedeceu aos seguintes princípios e ressalvas:")
    
    ressalvas = [
        "O vistoriador signatário inspecionou pessoalmente o objeto e o relatório técnico foi elaborado pelo próprio e ninguém, a não ser o mesmo, preparou as análises e as respectivas conclusões;",
        "O Relatório técnico foi elaborado com estrita observância dos postulados constantes do Código de Ética Profissional;",
        "Os honorários profissionais do signatário não estão, de qualquer forma, subordinados às conclusões deste relatório técnico;",
        "O vistoriador signatário não tem nenhuma inclinação pessoal em relação à matéria envolvida neste relatório técnico no presente, nem contempla para o futuro, qualquer interesse no bem objeto deste relatório técnico."
    ]
    
    for i, ressalva in enumerate(ressalvas, start=1):
        p = doc.add_paragraph()
        p.style = 'List Bullet'
        p.add_run(f"{chr(96+i)}) {ressalva}")
    
    # 2. OBJETIVO
    doc.add_page_break()
    doc.add_heading('OBJETIVO', level=1)
    
    p = doc.add_paragraph()
    p.add_run("O presente Laudo Técnico de Inspeção Predial foi solicitado pelo ")
    p.add_run(dados.get('contratante', '')).bold = True
    p.add_run(", CNPJ: ")
    p.add_run(dados.get('cnpj', '')).bold = True
    p.add_run(", elaborado pelo Engenheiro Civil, Silvio Augusto Barbosa de Albuquerque Filho, CREA-PE nº 054787D-PE")
    
    if dados.get('art_numero'):
        p.add_run(f", com registro da ART nº{dados['art_numero']} do presente documento.")
    else:
        p.add_run(".")
    
    doc.add_paragraph("A inspeção irá registrar as anomalias e falhas prediais por meio de um check-up da edificação.")
    
    # Breve relato
    if dados.get('breve_relato'):
        doc.add_heading('BREVE RELATO', level=2)
        p = doc.add_paragraph()
        p.add_run(f"Entre os dias {dados.get('dias_vistoria', '')} foram realizadas vistorias pela empresa ")
        p.add_run(dados.get('contratada', '')).bold = True
        p.add_run(" a pedido do ")
        p.add_run(dados.get('contratante', '')).bold = True
        p.add_run(" no imóvel localizado ")
        p.add_run(dados.get('endereco', '')).bold = True
        p.add_run(", no qual afirma:")
        
        # Processar breve relato
        doc.add_paragraph()
        relato_linhas = dados['breve_relato'].split('\n')
        for i, linha in enumerate(relato_linhas, start=1):
            if linha.strip():
                p = doc.add_paragraph()
                p.style = 'List Number'
                p.add_run(linha.strip())
    
    # 3. DESCRIÇÃO DO OBJETO
    doc.add_page_break()
    doc.add_heading('DESCRIÇÃO DO OBJETO INSPECIONADO', level=1)
    
    p = doc.add_paragraph()
    p.add_run(f"Trata-se de um empreendimento do tipo {dados.get('tipo_empreendimento', '')}, ")
    p.add_run(dados.get('info_localizacao', ''))
    p.add_run(f". O edifício está {'ocupado' if dados.get('ocupado') == 'Sim' else 'desocupado'}.")
    
    # 12. DOCUMENTAÇÕES
    doc.add_page_break()
    doc.add_heading('DOCUMENTAÇÕES SOLICITADAS E DOCUMENTAÇÕES DISPONIBILIZADAS:', level=1)
    
    docs_disponibilizadas = dados.get('docs_disponibilizadas', [])
    
    for doc_nome in DOCUMENTACOES:
        p = doc.add_paragraph()
        p.style = 'List Bullet'
        if doc_nome in docs_disponibilizadas:
            p.add_run(f"{doc_nome} - ")
            run = p.add_run("DISPONIBILIZADA")
            run.bold = True
        else:
            p.add_run(f"{doc_nome} - ")
            run = p.add_run("AUSENTE")
            run.bold = True
    
    if dados.get('obs_docs'):
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Obs: ").bold = True
        p.add_run(dados['obs_docs'])
    
    # 13. ANAMNESE
    doc.add_page_break()
    doc.add_heading('ANAMNESE', level=1)
    doc.add_paragraph(dados.get('anamnese', ''))
    
    doc.add_paragraph("A coordenação de dados se dá por meio de textos classificando as constatações de modo que as análises serão divididas de acordo com os arquivos anexos.")
    
    # Processar eventos
    for evento in eventos:
        doc.add_paragraph()
        
        # Título do evento
        p = doc.add_paragraph()
        run = p.add_run(f"EVENTO {evento['numero']:02d}: {evento['nome']}")
        run.bold = True
        run.font.size = Pt(12)
        
        # Informações do evento
        p = doc.add_paragraph()
        p.add_run("Localização: ").bold = True
        p.add_run(evento.get('localizacao', ''))
        
        p = doc.add_paragraph()
        p.add_run("Anomalia: ").bold = True
        p.add_run(", ".join(evento.get('anomalias', [])))
        
        p = doc.add_paragraph()
        p.add_run("Provável causa: ").bold = True
        p.add_run(evento.get('causa', ''))
        
        p = doc.add_paragraph()
        p.add_run("Consequência da anomalia: ").bold = True
        p.add_run(", ".join(evento.get('consequencias', [])))
        
        p = doc.add_paragraph()
        p.add_run("Patamar de urgência: ").bold = True
        p.add_run(evento.get('prioridade', ''))
        
        p = doc.add_paragraph()
        p.add_run("Uso: ").bold = True
        p.add_run(evento.get('uso', ''))
        
        p = doc.add_paragraph()
        p.add_run("Recomendação técnica: ").bold = True
        p.add_run(", ".join(evento.get('recomendacoes', [])))
    
    # Tabela resumo
    if eventos:
        doc.add_page_break()
        doc.add_heading('Resumo de Eventos por Prioridade', level=2)
        
        # Criar tabela
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Cabeçalho
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'EVENTO'
        hdr_cells[1].text = 'ANOMALIA'
        hdr_cells[2].text = 'PRIORIDADE'
        
        # Ordenar eventos por prioridade
        eventos_ordenados = sorted(eventos, key=lambda x: (x['prioridade'], x['numero']))
        
        # Adicionar linhas
        for evento in eventos_ordenados:
            row = table.add_row()
            row.cells[0].text = f"EVENTO {evento['numero']:02d}"
            row.cells[1].text = ", ".join(evento['anomalias'])
            row.cells[2].text = evento['prioridade'].split()[-1]
    
    # 14. LAUDO TÉCNICO
    doc.add_page_break()
    doc.add_heading('LAUDO TÉCNICO', level=1)
    
    if dados.get('texto_laudo'):
        doc.add_paragraph(dados['texto_laudo'])
    else:
        # Texto padrão
        texto_padrao = """O presente laudo técnico de inspeção predial foi elaborado com base nas vistorias realizadas entre os dias {} na edificação localizada em {}, pertencente ao {}. O objetivo foi avaliar as condições gerais da edificação, com foco na integridade estrutural, funcionalidade dos sistemas construtivos, segurança dos usuários, e condições de habitabilidade, em conformidade com as diretrizes da ABNT NBR 16747:2020 e da NBR 13752:2024."""
        
        doc.add_paragraph(texto_padrao.format(
            dados.get('dias_vistoria', ''),
            dados.get('endereco', ''),
            dados.get('contratante', '')
        ))
    
    # 15. DATA DO RELATÓRIO
    doc.add_page_break()
    doc.add_heading('DATA DO RELATÓRIO TÉCNICO', level=1)
    
    p = doc.add_paragraph()
    p.add_run(f"Em {dados.get('data_laudo').strftime('%d de %B de %Y')}, ")
    p.add_run("com base nos trabalhos aqui representados encerramos o presente relatório técnico.")
    
    # Assinatura
    for _ in range(3):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("_" * 50)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Eng. Responsável: Eng. Silvio Albuquerque Filho").bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("CREA: 054787D-PE").bold = True
    
    if dados.get('art_numero'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"ART: {dados['art_numero']}").bold = True
    
    return doc

# Interface principal
st.title("🏢 Gerador de Laudos de Inspeção Predial")

# Sidebar
with st.sidebar:
    st.title("📋 Menu")
    st.info("Sistema profissional para geração de laudos técnicos")
    
    if st.button("🆕 Novo Laudo"):
        st.session_state.dados_laudo = {}
        st.session_state.eventos = []
    
    if st.button("💾 Salvar Rascunho"):
        if st.session_state.dados_laudo:
            nome = f"Rascunho_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            st.session_state.laudos_salvos[nome] = {
                'dados': st.session_state.dados_laudo.copy(),
                'eventos': st.session_state.eventos.copy()
            }
            st.success(f"Salvo: {nome}")
    
    if st.session_state.laudos_salvos:
        st.divider()
        st.subheader("📂 Laudos Salvos")
        for nome in st.session_state.laudos_salvos:
            if st.button(f"📄 {nome}", key=nome):
                st.session_state.dados_laudo = st.session_state.laudos_salvos[nome]['dados']
                st.session_state.eventos = st.session_state.laudos_salvos[nome]['eventos']
                st.success(f"Carregado: {nome}")

# Tabs principais
tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "📝 Dados Básicos",
    "📍 Localização",
    "📋 Documentação",
    "🔍 Eventos",
    "📄 Gerar Laudo"
])

# TAB 1 - DADOS BÁSICOS
with tab1:
    st.subheader("Informações Básicas do Laudo")
    
    col1, col2 = st.columns(2)
    
    with col1:
        contratante = st.text_input(
            "Nome do Contratante*",
            value=st.session_state.dados_laudo.get('contratante', ''),
            help="Ex: Ser Educacional S.A - Centro Universitário"
        )
        
        cnpj = st.text_input(
            "CNPJ/CPF*",
            value=st.session_state.dados_laudo.get('cnpj', ''),
            help="Formato: XX.XXX.XXX/XXXX-XX"
        )
        
        data_laudo = st.date_input(
            "Data do Laudo*",
            value=st.session_state.dados_laudo.get('data_laudo', datetime.now())
        )
        
        contratada_opcao = st.selectbox(
            "Empresa Contratada*",
            options=OPCOES['contratada']
        )
        
        if contratada_opcao == "Outra":
            contratada = st.text_input("Nome da Contratada")
        else:
            contratada = contratada_opcao
    
    with col2:
        dias_vistoria = st.text_input(
            "Dias de Vistoria*",
            value=st.session_state.dados_laudo.get('dias_vistoria', ''),
            help="Ex: 08 a 11/07/2025"
        )
        
        art_numero = st.text_input(
            "Número da ART",
            value=st.session_state.dados_laudo.get('art_numero', ''),
            help="Deixe em branco se não houver"
        )
        
        cidade_estado = st.text_input(
            "Cidade-Estado*",
            value=st.session_state.dados_laudo.get('cidade_estado', ''),
            help="Ex: Natal-RN"
        )
        
        ocupado = st.radio(
            "O empreendimento está ocupado?",
            ["Sim", "Não"],
            horizontal=True
        )
    
    # Salvar dados
    st.session_state.dados_laudo.update({
        'contratante': contratante,
        'cnpj': cnpj,
        'data_laudo': data_laudo,
        'contratada': contratada,
        'dias_vistoria': dias_vistoria,
        'art_numero': art_numero,
        'cidade_estado': cidade_estado,
        'ocupado': ocupado
    })

# TAB 2 - LOCALIZAÇÃO
with tab2:
    st.subheader("📍 Localização do Imóvel")
    
    endereco = st.text_area(
        "Endereço Completo*",
        value=st.session_state.dados_laudo.get('endereco', ''),
        height=100,
        help="Digite o endereço completo incluindo CEP"
    )
    
    col1, col2 = st.columns(2)
    
    with col1:
        tipo_opcao = st.selectbox(
            "Tipo de Empreendimento*",
            options=OPCOES['tipo_empreendimento']
        )
        if tipo_opcao == "Outro":
            tipo_empreendimento = st.text_input("Especifique o tipo")
        else:
            tipo_empreendimento = tipo_opcao
    
    with col2:
        info_localizacao = st.text_area(
            "Informações sobre a Localização",
            value="encontra-se em área urbanizada, perto de comércio e com estrutura desenvolvida de saneamento básico",
            height=100
        )
    
    st.session_state.dados_laudo.update({
        'endereco': endereco,
        'tipo_empreendimento': tipo_empreendimento,
        'info_localizacao': info_localizacao
    })

# TAB 3 - DOCUMENTAÇÃO
with tab3:
    st.subheader("📋 Documentações")
    st.info("Selecione as documentações que foram disponibilizadas")
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("✅ Marcar Todas"):
            st.session_state['todas_docs'] = True
            st.rerun()
    with col2:
        if st.button("❌ Desmarcar Todas"):
            st.session_state['todas_docs'] = False
            st.rerun()
    
    st.divider()
    
    docs_disponibilizadas = []
    for i, doc in enumerate(DOCUMENTACOES):
        valor = st.session_state.get('todas_docs', False)
        if st.checkbox(doc, value=valor, key=f"doc_{i}"):
            docs_disponibilizadas.append(doc)
    
    obs_docs = st.text_area(
        "Observações sobre documentações",
        value=st.session_state.dados_laudo.get('obs_docs', ''),
        help="Ex: Obs: Das documentações solicitadas apenas os projetos arquitetônicos..."
    )
    
    st.session_state.dados_laudo.update({
        'docs_disponibilizadas': docs_disponibilizadas,
        'obs_docs': obs_docs
    })

# TAB 4 - EVENTOS
with tab4:
    st.subheader("🔍 Eventos de Inspeção")
    
    # Breve relato
    st.subheader("Breve Relato")
    breve_relato = st.text_area(
        "Digite o breve relato da contratante (cada linha será numerada)",
        value=st.session_state.dados_laudo.get('breve_relato', ''),
        height=200,
        help="Ex: Ocupam o imóvel há 2 anos\nNão possuem Manual de Uso..."
    )
    st.session_state.dados_laudo['breve_relato'] = breve_relato
    
    st.divider()
    
    # Anamnese
    anamnese = st.text_area(
        "Anamnese",
        value=st.session_state.dados_laudo.get('anamnese', 
            "Os usuários da edificação pontuam de forma simplificada que perceberam uma deterioração comumente natural dos materiais componentes da edificação que estão em desconformidades, que por consequência está ocorrendo na edificação, incidências de infiltrações e problemas nas instalações elétricas e hidrossanitários, chegando à solicitação do presente laudo de inspeção."),
        height=150
    )
    st.session_state.dados_laudo['anamnese'] = anamnese
    
    st.divider()
    
    # Gerenciamento de Eventos
    col1, col2, col3 = st.columns([2, 1, 1])
    with col1:
        st.subheader(f"Total de Eventos: {len(st.session_state.eventos)}")
    with col2:
        if st.button("➕ Adicionar Evento"):
            novo_evento = {
                'numero': len(st.session_state.eventos) + 1,
                'nome': '',
                'localizacao': 'Generalidades',
                'anomalias': [],
                'causa': 'Funcional',
                'consequencias': [],
                'prioridade': 'Prioridade 2',
                'uso': 'Regular',
                'recomendacoes': [],
                'imagens': []
            }
            st.session_state.eventos.append(novo_evento)
            st.rerun()
    with col3:
        if st.button("🗑️ Limpar Todos"):
            st.session_state.eventos = []
            st.rerun()
    
    # Exibir eventos
    for idx, evento in enumerate(st.session_state.eventos):
        with st.expander(f"📌 EVENTO {evento['numero']:02d}: {evento.get('nome', 'Sem nome')}", expanded=True):
            col1, col2 = st.columns([3, 1])
            
            with col1:
                evento['nome'] = st.text_input(
                    "Nome do Evento",
                    value=evento.get('nome', ''),
                    key=f"nome_{idx}"
                )
            
            with col2:
                if st.button("❌ Remover", key=f"remove_{idx}"):
                    st.session_state.eventos.pop(idx)
                   # Renumerar eventos
                   for i, evt in enumerate(st.session_state.eventos):
                       evt['numero'] = i + 1
                   st.rerun()
           
           # Localização
           loc_col1, loc_col2 = st.columns(2)
           with loc_col1:
               if st.checkbox("Generalidades", 
                             value=evento.get('localizacao') == 'Generalidades', 
                             key=f"gen_{idx}"):
                   evento['localizacao'] = "Generalidades"
           with loc_col2:
               loc_custom = st.text_input("Ou especifique:", key=f"loc_{idx}")
               if loc_custom:
                   evento['localizacao'] = loc_custom
           
           # Anomalias
           evento['anomalias'] = st.multiselect(
               "Anomalias",
               options=OPCOES['anomalias'],
               default=evento.get('anomalias', []),
               key=f"anom_{idx}"
           )
           
           # Causa
           evento['causa'] = st.selectbox(
               "Provável Causa",
               options=OPCOES['causas'],
               index=OPCOES['causas'].index(evento.get('causa', 'Funcional')),
               key=f"causa_{idx}"
           )
           
           # Consequências
           evento['consequencias'] = st.multiselect(
               "Consequências",
               options=OPCOES['consequencias'],
               default=evento.get('consequencias', []),
               key=f"cons_{idx}"
           )
           
           # Prioridade
           evento['prioridade'] = st.radio(
               "Patamar de Urgência",
               ["Prioridade 1", "Prioridade 2", "Prioridade 3"],
               index=["Prioridade 1", "Prioridade 2", "Prioridade 3"].index(
                   evento.get('prioridade', 'Prioridade 2')),
               key=f"prio_{idx}",
               horizontal=True
           )
           
           # Uso
           evento['uso'] = st.radio(
               "Uso",
               ["Regular", "Irregular"],
               index=["Regular", "Irregular"].index(evento.get('uso', 'Regular')),
               key=f"uso_{idx}",
               horizontal=True
           )
           
           # Recomendações
           evento['recomendacoes'] = st.multiselect(
               "Recomendações Técnicas",
               options=OPCOES['recomendacoes'],
               default=evento.get('recomendacoes', []),
               key=f"rec_{idx}"
           )
           
           # Upload de imagens
           st.write("📷 Imagens do Evento (2-3 imagens)")
           imgs = st.file_uploader(
               "Selecione as imagens",
               type=['png', 'jpg', 'jpeg'],
               accept_multiple_files=True,
               key=f"imgs_{idx}"
           )
           if imgs:
               if len(imgs) > 3:
                   st.warning("Máximo de 3 imagens. Usando apenas as 3 primeiras.")
                   evento['imagens'] = imgs[:3]
               else:
                   evento['imagens'] = imgs

# TAB 5 - GERAR LAUDO
with tab5:
   st.subheader("📄 Geração do Laudo Final")
   
   # Opções do texto
   opcao_texto = st.radio(
       "Como deseja gerar o texto do laudo?",
       ["📝 Usar texto padrão", "✏️ Escrever manualmente"]
   )
   
   texto_laudo = ""
   if opcao_texto == "📝 Usar texto padrão":
       texto_laudo = """O presente laudo técnico de inspeção predial foi elaborado com base nas vistorias realizadas entre os dias [DIAS], na edificação localizada na [ENDEREÇO], pertencente ao [CONTRATANTE]. O objetivo foi avaliar as condições gerais da edificação, com foco na integridade estrutural, funcionalidade dos sistemas construtivos, segurança dos usuários, e condições de habitabilidade, em conformidade com as diretrizes da ABNT NBR 16747:2020 e da NBR 13752:2024.

Com base na avaliação técnica criteriosa realizada nesta inspeção predial, conclui-se que a edificação objeto deste laudo apresenta um quadro patológico de natureza multifatorial, cujas manifestações indicam um nível de criticidade classificado como alto, com predominância de anomalias do tipo endógeno e funcional.

A avaliação sensorial in loco, realizada conforme os preceitos estabelecidos pela ABNT NBR 16747:2020 e demais normativas correlatas, evidenciou a presença de falhas recorrentes em sistemas de impermeabilização, revestimentos, esquadrias, pisos e elementos de acessibilidade, comprometendo a durabilidade, a funcionalidade e, em determinadas circunstâncias, a segurança e o conforto dos usuários da edificação.

Importa salientar que devido a idade da construção de mais de uma década e a ausência de um plano sistematizado de manutenção preventiva, bem como de documentação técnica incompleta, incluindo manuais de uso e operação, tem potencializado o surgimento e agravamento das patologias observadas. A inexistência de determinadas licenças legais e o uso indevido de determinados espaços reforçam a necessidade de regularização junto aos órgãos competentes.

Recomenda-se, com o devido grau de urgência e priorização, a execução das intervenções corretivas indicadas neste relatório, por meio da contratação de empresas especializadas, com responsabilidade técnica devidamente atribuída, a fim de assegurar a conformidade técnica, o atendimento aos requisitos normativos e a reabilitação plena dos sistemas construtivos comprometidos."""
       
       # Preview
       st.text_area("Preview do texto padrão", texto_laudo, height=300, disabled=True)
       
   else:
       texto_laudo = st.text_area(
           "Digite o texto completo do laudo",
           value=st.session_state.dados_laudo.get('texto_laudo', ''),
           height=400,
           help="Digite aqui o texto completo do laudo técnico"
       )
   
   st.session_state.dados_laudo['texto_laudo'] = texto_laudo
   
   # Opções finais
   st.divider()
   
   col1, col2 = st.columns(2)
   with col1:
       incluir_rodape = st.checkbox("Incluir rodapé", value=True)
       incluir_numeracao = st.checkbox("Incluir numeração de páginas", value=True)
   with col2:
       versao = st.number_input("Versão do documento", min_value=1, value=1)
   
   # Validação
   campos_obrigatorios = ['contratante', 'cnpj', 'endereco', 'cidade_estado', 'dias_vistoria']
   todos_preenchidos = all(st.session_state.dados_laudo.get(campo) for campo in campos_obrigatorios)
   
   if not todos_preenchidos:
       st.warning("⚠️ Preencha todos os campos obrigatórios antes de gerar o laudo")
       campos_faltando = [campo for campo in campos_obrigatorios if not st.session_state.dados_laudo.get(campo)]
       st.error(f"Campos faltando: {', '.join(campos_faltando)}")
   elif len(st.session_state.eventos) == 0:
       st.warning("⚠️ Adicione pelo menos um evento antes de gerar o laudo")
   else:
       col1, col2, col3 = st.columns([1, 2, 1])
       with col2:
           if st.button("🚀 GERAR LAUDO COMPLETO", type="primary", use_container_width=True):
               with st.spinner("Gerando documento... Por favor aguarde..."):
                   try:
                       # Gerar documento
                       doc = gerar_documento_completo(
                           st.session_state.dados_laudo,
                           st.session_state.eventos,
                           incluir_rodape,
                           incluir_numeracao,
                           versao
                       )
                       
                       # Salvar em buffer
                       doc_buffer = io.BytesIO()
                       doc.save(doc_buffer)
                       doc_buffer.seek(0)
                       
                       # Nome do arquivo
                       contratante_nome = st.session_state.dados_laudo['contratante'].replace(' ', '_')
                       data_str = st.session_state.dados_laudo['data_laudo'].strftime('%Y%m%d')
                       nome_arquivo = f"LAUDO_{contratante_nome}_{data_str}_v{versao}.docx"
                       
                       # Sucesso e download
                       st.success(f"✅ Laudo gerado com sucesso!")
                       st.balloons()
                       
                       # Botão de download
                       col1, col2, col3 = st.columns([1, 2, 1])
                       with col2:
                           st.download_button(
                               label="📥 BAIXAR LAUDO",
                               data=doc_buffer,
                               file_name=nome_arquivo,
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                               use_container_width=True
                           )
                       
                       # Salvar na memória
                       st.session_state.laudos_salvos[nome_arquivo] = {
                           'dados': st.session_state.dados_laudo.copy(),
                           'eventos': st.session_state.eventos.copy(),
                           'data_criacao': datetime.now()
                       }
                       
                       st.info(f"📁 Arquivo: {nome_arquivo}")
                       
                       # Informações do laudo
                       with st.expander("📊 Resumo do Laudo Gerado"):
                           col1, col2 = st.columns(2)
                           with col1:
                               st.write("**Contratante:**", st.session_state.dados_laudo['contratante'])
                               st.write("**CNPJ:**", st.session_state.dados_laudo['cnpj'])
                               st.write("**Endereço:**", st.session_state.dados_laudo['endereco'])
                           with col2:
                               st.write("**Total de Eventos:**", len(st.session_state.eventos))
                               st.write("**Prioridade 1:**", sum(1 for e in st.session_state.eventos if e['prioridade'] == 'Prioridade 1'))
                               st.write("**Prioridade 2:**", sum(1 for e in st.session_state.eventos if e['prioridade'] == 'Prioridade 2'))
                               st.write("**Prioridade 3:**", sum(1 for e in st.session_state.eventos if e['prioridade'] == 'Prioridade 3'))
                       
                   except Exception as e:
                       st.error(f"❌ Erro ao gerar documento: {str(e)}")
                       st.error("Por favor, verifique se todos os campos estão preenchidos corretamente.")

# Footer
st.divider()
st.markdown("""
<div style='text-align: center; color: gray; padding: 20px;'>
   <p>Sistema de Geração de Laudos de Inspeção Predial v1.0</p>
   <p>Desenvolvido para facilitar a criação de laudos técnicos profissionais</p>
</div>
""", unsafe_allow_html=True)